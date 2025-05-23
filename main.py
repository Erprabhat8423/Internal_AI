from fastapi import FastAPI, File, UploadFile, Query
from sentence_transformers import SentenceTransformer
import fitz  # PyMuPDF for PDF
import docx
import faiss
import os
import groq
import numpy as np
from sqlalchemy import create_engine, Column, Integer, String, Text
from sqlalchemy.orm import sessionmaker, declarative_base
import dotenv
from sqlalchemy.orm import Session
from os import getenv
from typing import List, Dict, Optional

from pydantic import BaseModel
from fastapi import Body
from fastapi import HTTPException


dotenv.load_dotenv()
GROQ_API_KEY = getenv("GROQ_API_KEY")

# Initialize FastAPI app
app = FastAPI()

# Initialize Groq Client
groq_client = groq.Client(api_key=GROQ_API_KEY)

# Load embedding model
model = SentenceTransformer("all-MiniLM-L6-v2")  # 384-dimension

FAISS_INDEX_PATH = "faiss_index.bin"
embedding_dim = 384
index = faiss.IndexFlatL2(embedding_dim)

# Database setup
DATABASE_URL = "mysql+pymysql://root:12345678@localhost/practice_document_rag_qa"
engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

class QueryPayload(BaseModel):
    question: str
    context: Optional[str] = ""

# SQLAlchemy Document model
class Document(Base):
    __tablename__ = "documents"
    id = Column(Integer, primary_key=True, index=True)
    filename = Column(String, nullable=False)
    content = Column(Text, nullable=False)
    faiss_index = Column(Integer, unique=True, nullable=True)

Base.metadata.create_all(bind=engine)

# Utility functions
def extract_text_from_pdf(pdf_file):
    doc = fitz.open(stream=pdf_file, filetype="pdf")
    return "\n".join(page.get_text("text") for page in doc).strip()

def extract_text_from_docx(docx_file):
    doc = docx.Document(docx_file)
    return "\n".join(para.text for para in doc.paragraphs).strip()

def get_embedding(text):
    embedding = model.encode(text)
    return embedding / np.linalg.norm(embedding)

def save_faiss_index():
    faiss.write_index(index, FAISS_INDEX_PATH)

def load_faiss_index():
    global index
    if os.path.exists(FAISS_INDEX_PATH):
        index = faiss.read_index(FAISS_INDEX_PATH)
        print("FAISS index loaded with", index.ntotal, "entries")
    else:
        print("No FAISS index found, using fresh one.")

# Load index at startup
load_faiss_index()

# Upload document API
@app.post("/upload/")
async def upload_document(file: UploadFile = File(...)):
    content = ""
    # ✅ Check if file already exists in the database
    db = SessionLocal()
    existing_doc = db.query(Document).filter_by(filename=file.filename).first()
    if existing_doc:
        db.close()
        return {"error": f"File '{file.filename}' already exists."}


    if file.filename.endswith(".pdf"):
        content = extract_text_from_pdf(await file.read())
    elif file.filename.endswith(".docx"):
        content = extract_text_from_docx(file.file)
    else:
        return {"error": "Unsupported file format"}

    embedding = get_embedding(content)
    load_faiss_index()
    index.add(np.array([embedding], dtype=np.float32))
    faiss_index = index.ntotal - 1
    save_faiss_index()

    db = SessionLocal()
    new_doc = Document(filename=file.filename, content=content, faiss_index=faiss_index)
    db.add(new_doc)
    db.commit()
    db.close()

    return {
        "message": "Document uploaded and embedded successfully",
        "faiss_index": faiss_index
    }

def extract_list_items(text: str):
    lines = text.strip().split("\n")
    return [line.lstrip("•-123. ").strip() for line in lines if line.strip()]

# Query documents
@app.post("/query/")
async def query_document(payload: QueryPayload):
    question = payload.question
    context = payload.context or ""
    query_embedding = get_embedding(question)
    k = 3
    distances, indices = index.search(np.array([query_embedding], dtype=np.float32), k)
    matched_indices = [int(idx) for idx in indices[0] if idx != -1]

    if not matched_indices:
        return {"answer": "I couldn't find anything useful in the documents.", "source": None}

    db: Session = SessionLocal()
    matched_docs = db.query(Document).filter(Document.faiss_index.in_(matched_indices)).all()
    db.close()

    if not matched_docs:
        return {"answer": "No document found for your query.", "source": None}

    def truncate_text(text, max_chars=2000):
        return text[:max_chars] if text else ""

    combined_content = "\n\n".join([
        f"{doc.filename}:\n{truncate_text(doc.content)}"
        for doc in matched_docs
    ])

    top_doc_name = matched_docs[0].filename if matched_docs else None

    strict_prompt = f"""
    You are a helpful and smart assistant.

    Use the DOCUMENTS below and the prior CHAT CONTEXT to answer the QUESTION.
    If the document has the answer — prioritize it. If not, use your reasoning based on the conversation to respond.

    DOCUMENTS:
    {combined_content}

    CHAT CONTEXT:
    {context}

    QUESTION:
    {question}

    Respond concisely and clearly.
    """


    try:
        response = groq_client.chat.completions.create(
            model="deepseek-r1-distill-llama-70b",
            messages=[
                {"role": "system", "content": "You are a strict document-based QA assistant."},
                {"role": "user", "content": strict_prompt}
            ],
            temperature=0.3,
            max_tokens=1000
        )
        exact_answer = response.choices[0].message.content.strip()
    except Exception as e:
        return {"answer": f"Unable to connect with model.", "source": None}

    fallback_phrases = [
        "no mention", "does not appear", "no information", "not contain any information",
        "there is no", "unable to find", "not found in the provided documents"
    ]
    if any(phrase in exact_answer.lower() for phrase in fallback_phrases):
        return {
            "answer": "I'm sorry, I couldn't find an exact answer in my knowledge base.",
            "source": None
        }
    return {"answer": exact_answer, "source": top_doc_name}